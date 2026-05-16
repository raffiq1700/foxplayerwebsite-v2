from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_styled_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0x0A, 0x11, 0x28) # Dark Blue
    if level == 1:
        run.font.size = Pt(24)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37) # Gold
    return heading

def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "[LOGO PLACEHOLDER] FoxPlayer Algo Technologies | Company Profile & Service Portfolio 2026"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.runs[0].font.size = Pt(9)
        header_para.runs[0].font.color.rgb = RGBColor(0x8B, 0x95, 0xA5)
        
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "www.foxplayer.co.in | raffiq_sr@yahoo.co.in"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)
        footer_para.runs[0].font.color.rgb = RGBColor(0x8B, 0x95, 0xA5)

doc = Document()

# Cover Page
doc.add_paragraph()
logo_para = doc.add_paragraph("[COMPANY LOGO PLACEHOLDER]")
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_para.runs[0].font.size = Pt(14)

add_styled_heading(doc, 'FoxPlayer Algo Technologies', 1)

tagline = doc.add_paragraph('Building Institutional-Grade Trading Platforms, Software Solutions, and Digital Products.')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].font.size = Pt(14)
tagline.runs[0].font.italic = True
tagline.runs[0].font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)

doc.add_paragraph()
subtitle = doc.add_paragraph('Company Profile & Service Portfolio 2026')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.bold = True

for _ in range(5):
    doc.add_paragraph()

contact_para = doc.add_paragraph("Website: https://www.foxplayer.co.in\nEmail: raffiq_sr@yahoo.co.in\nLinkedIn: linkedin.com/in/mohamed-raffiq-6b97a6a7")
contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_para.runs[0].font.size = Pt(11)

doc.add_page_break()

# Page 2: About Us
add_styled_heading(doc, 'About FoxPlayer Algo Technologies', 2)
doc.add_paragraph("FoxPlayer Algo Technologies helps traders, brokers, startups, and businesses transform ideas into scalable software products. We combine deep expertise in algorithmic trading, broker APIs, financial markets, and full-stack development to build powerful, secure, and production-ready solutions tailored to each client's goals.")
doc.add_paragraph("FoxPlayer Algo Technologies is a specialized software development company focused on algorithmic trading platforms, broker API integrations, backtesting engines, custom web applications, mobile apps, SaaS products, and enterprise software solutions.")

doc.add_heading('Mission', level=3)
doc.add_paragraph("To deliver world-class software solutions that empower businesses and traders through automation, innovation, and cutting-edge technology.")

doc.add_heading('Vision', level=3)
doc.add_paragraph("To become a globally trusted technology partner in algorithmic trading, FinTech, and custom software development.")

doc.add_page_break()

# Page 3: Core Services
add_styled_heading(doc, 'Core Services', 2)
services = [
    "Algorithmic Trading Software Development",
    "Broker API Integration (Shoonya, Zerodha, Alice Blue, Angel One)",
    "White Label Trading Platforms",
    "Backtesting and Strategy Automation Engines",
    "TradingView Webhook Automation",
    "Custom Web Development",
    "Mobile App Development (Android and iOS)",
    "SaaS Product Development",
    "E-Commerce Solutions",
    "UI/UX Design and Dashboard Development"
]
for service in services:
    doc.add_paragraph(service, style='List Bullet')

doc.add_page_break()

# Page 4: Why Choose Us
add_styled_heading(doc, 'Why Choose Us', 2)
reasons = [
    "Specialized expertise in Indian financial markets",
    "Deep knowledge of broker APIs and trading automation",
    "High-performance applications using Python, C++, Next.js, and PostgreSQL",
    "SEO-optimized websites and scalable SaaS products",
    "Secure, maintainable, and enterprise-grade code",
    "End-to-end development from concept to deployment",
    "Transparent communication and dedicated support"
]
for reason in reasons:
    doc.add_paragraph(reason, style='List Bullet')

doc.add_page_break()

# Page 5: Technology Stack
add_styled_heading(doc, 'Technology Stack', 2)
tech_stack = {
    "Backend": "Python, FastAPI, Flask, Django, C++",
    "Frontend": "Next.js, React, Tailwind CSS, TypeScript",
    "Database": "PostgreSQL, SQLite",
    "Infrastructure": "Docker, Linux, AWS, VPS Hosting",
    "Trading Integrations": "Shoonya, Zerodha Kite Connect, Alice Blue, Angel One, TradingView Webhooks"
}
for category, techs in tech_stack.items():
    p = doc.add_paragraph()
    p.add_run(f"{category}: ").bold = True
    p.add_run(techs)

doc.add_page_break()

# Page 6: Industries We Serve
add_styled_heading(doc, 'Industries We Serve', 2)
industries = [
    "Brokerage Firms",
    "Proprietary Trading Firms",
    "Asset Management Companies",
    "FinTech Startups",
    "E-Commerce Businesses",
    "Educational Platforms",
    "Enterprise Clients"
]
for industry in industries:
    doc.add_paragraph(industry, style='List Bullet')

doc.add_page_break()

# Page 7: Development Process
add_styled_heading(doc, 'Development Process', 2)
process = [
    "Requirement Analysis",
    "Solution Architecture",
    "UI/UX Design",
    "Agile Development",
    "Testing & Quality Assurance",
    "Deployment",
    "Ongoing Maintenance & Support"
]
for i, step in enumerate(process, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_page_break()

# Page 8: Trusted Market Integrations
add_styled_heading(doc, 'Trusted Market Integrations', 2)
doc.add_paragraph("Integrations with leading platforms including:")
integrations = ["Shoonya", "Zerodha", "Alice Blue", "Angel One", "NSE", "MCX", "SEBI", "Groww"]
for integration in integrations:
    doc.add_paragraph(integration, style='List Bullet')

doc.add_page_break()

# Page 9: Contact Information
add_styled_heading(doc, 'Contact Information', 2)
doc.add_paragraph("Founder: Mohamed Raffiq\nWebsite: https://www.foxplayer.co.in\nEmail: raffiq_sr@yahoo.co.in\nLinkedIn: https://www.linkedin.com/in/mohamed-raffiq-6b97a6a7")

doc.add_paragraph()
cta = doc.add_paragraph("Ready to build your next trading platform, SaaS product, or enterprise application?\nContact FoxPlayer Algo Technologies today to discuss your project.")
cta.alignment = WD_ALIGN_PARAGRAPH.CENTER
cta.runs[0].bold = True
cta.runs[0].font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)

add_header_footer(doc)

doc.save('FoxPlayer_Algo_Technologies_Brochure.docx')
print('DOCX generated successfully')
